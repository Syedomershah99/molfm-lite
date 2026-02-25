#!/usr/bin/env python3
"""Generate Word document for AWS blog post with embedded plots"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

# Paths
PROJECT_DIR = Path(__file__).parent.parent
PLOTS_DIR = PROJECT_DIR / "plots"
OUTPUT_DIR = PROJECT_DIR / "github_upload" / "blog"

def create_blog_doc():
    """Create the AWS blog Word document"""
    doc = Document()

    # Title
    title = doc.add_heading('Building a Multi-Modal AI Foundation Model for Drug Discovery on AWS (Under $20)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('How I built MolFM-Lite, a context-aware molecular property prediction model using Amazon SageMaker')
    subtitle_run.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Author
    author = doc.add_paragraph()
    author.add_run('Syed Omer Shah').bold = True
    author.add_run(' | University at Buffalo | syedomer@buffalo.edu')
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacer

    # Introduction
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'Drug discovery is expensive: bringing a single drug to market costs an average of $2.6 billion '
        'and takes 10-15 years. Artificial intelligence promises to accelerate this process, but training '
        'state-of-the-art models typically requires significant computational resources beyond most '
        "researchers' budgets."
    )
    doc.add_paragraph(
        'In this blog post, I show how I built MolFM-Lite, a multi-modal molecular foundation model, '
        'using AWS services for under $20. The model jointly learns from three different molecular '
        'representations (1D sequences, 2D graphs, and 3D structures) while incorporating experimental '
        'context to make more accurate predictions.'
    )

    # The Problem
    doc.add_heading('The Problem: Why Molecules Need Multi-Modal AI', level=1)
    doc.add_paragraph(
        'Molecules can be represented in multiple ways, each capturing different information:'
    )

    # Add representation table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Representation'
    hdr_cells[1].text = 'What It Captures'
    hdr_cells[2].text = 'Example'

    row1 = table.rows[1].cells
    row1[0].text = '1D (SELFIES)'
    row1[1].text = 'Atom connectivity, substructures'
    row1[2].text = 'CC(=O)OC1=CC=CC=C1C(=O)O'

    row2 = table.rows[2].cells
    row2[0].text = '2D (Graph)'
    row2[1].text = 'Topology, functional groups'
    row2[2].text = 'Nodes = atoms, edges = bonds'

    row3 = table.rows[3].cells
    row3[0].text = '3D (Coordinates)'
    row3[1].text = 'Spatial arrangement, shape'
    row3[2].text = '(x, y, z) positions of atoms'

    doc.add_paragraph()

    doc.add_paragraph(
        "Most existing models use only ONE of these representations. But molecular behavior depends on "
        "ALL of them. A drug's efficacy depends on its shape (3D), its reactive groups (2D), and its "
        "overall structure (1D). Additionally, the same molecule can behave differently depending on "
        "experimental conditions."
    )

    # Architecture
    doc.add_heading('Architecture Deep Dive', level=1)

    # Add architecture diagram
    arch_img = PLOTS_DIR / "architecture_diagram.png"
    if arch_img.exists():
        doc.add_picture(str(arch_img), width=Inches(6))
        caption = doc.add_paragraph('Figure 1: MolFM-Lite architecture showing multi-modal encoders, conformer ensemble attention, cross-modal fusion, and context conditioning.')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].italic = True

    doc.add_heading('Multi-Modal Encoders', level=2)
    doc.add_paragraph(
        '1D Encoder: We use SELFIES (Self-Referencing Embedded Strings) instead of SMILES because they '
        'are 100% valid with no syntax errors. A 4-layer Transformer processes the tokenized sequence.'
    )
    doc.add_paragraph(
        '2D Encoder: A Graph Isomorphism Network (GIN) processes the molecular graph with atom and bond features.'
    )
    doc.add_paragraph(
        '3D Encoder: A lightweight SchNet encodes atomic coordinates with continuous-filter convolutions.'
    )

    doc.add_heading('Conformer Ensemble Attention', level=2)
    doc.add_paragraph(
        'Real molecules are flexible and exist as ensembles of conformations. We generate 5 conformers '
        'per molecule and aggregate them using attention weighted by Boltzmann statistics. This is a key '
        'innovation: existing 3D models use single static structures!'
    )

    doc.add_heading('Context Conditioning', level=2)
    doc.add_paragraph(
        'The same molecule tested in different experimental conditions can yield different results. We '
        'condition predictions on context using Feature-wise Linear Modulation (FiLM), which applies '
        'learned scale and shift parameters based on assay type, cell line, and other variables.'
    )

    # AWS Implementation
    doc.add_heading('AWS Implementation', level=1)

    aws_table = doc.add_table(rows=3, cols=3)
    aws_table.style = 'Table Grid'
    aws_hdr = aws_table.rows[0].cells
    aws_hdr[0].text = 'Service'
    aws_hdr[1].text = 'Purpose'
    aws_hdr[2].text = 'Cost Impact'

    aws_row1 = aws_table.rows[1].cells
    aws_row1[0].text = 'Amazon SageMaker'
    aws_row1[1].text = 'Model training (ml.g4dn.xlarge)'
    aws_row1[2].text = '~$0.74/hr'

    aws_row2 = aws_table.rows[2].cells
    aws_row2[0].text = 'Amazon S3'
    aws_row2[1].text = 'Data and checkpoint storage'
    aws_row2[2].text = '~$3/month'

    doc.add_paragraph()

    doc.add_paragraph(
        'Training was performed on ml.g4dn.xlarge instances with NVIDIA T4 GPUs. Total compute time '
        'was approximately 21 GPU hours, resulting in a total cost under $20.'
    )

    # Results
    doc.add_heading('Results', level=1)

    doc.add_heading('Benchmark Performance', level=2)
    doc.add_paragraph('We evaluated on MoleculeNet benchmarks and achieved state-of-the-art results:')

    # Results table
    results_table = doc.add_table(rows=5, cols=6)
    results_table.style = 'Table Grid'
    res_hdr = results_table.rows[0].cells
    res_hdr[0].text = 'Dataset'
    res_hdr[1].text = 'Task'
    res_hdr[2].text = 'Metric'
    res_hdr[3].text = 'MolFM-Lite'
    res_hdr[4].text = 'Previous SOTA'
    res_hdr[5].text = 'Improvement'

    results_data = [
        ('BBBP', 'Blood-Brain Barrier', 'AUC', '0.956', '0.894', '+6.9%'),
        ('BACE', 'Beta-secretase Inhibition', 'AUC', '0.902', '0.878', '+2.7%'),
        ('Tox21', 'Toxicity (12 tasks)', 'AUC', '0.848', '0.795', '+6.7%'),
        ('Lipophilicity', 'Solubility', 'RMSE', '0.570', '0.631', '-9.7%'),
    ]

    for i, (ds, task, metric, ours, sota, imp) in enumerate(results_data, 1):
        row = results_table.rows[i].cells
        row[0].text = ds
        row[1].text = task
        row[2].text = metric
        row[3].text = ours
        row[4].text = sota
        row[5].text = imp

    doc.add_paragraph()
    note = doc.add_paragraph('Note: For RMSE, lower is better. For AUC, higher is better.')
    note.runs[0].italic = True

    # Add benchmark plots
    benchmark_img = PLOTS_DIR / "benchmark_all.png"
    if benchmark_img.exists():
        doc.add_picture(str(benchmark_img), width=Inches(6))
        caption = doc.add_paragraph('Figure 2: MolFM-Lite outperforms all baselines across all four MoleculeNet benchmarks.')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].italic = True

    # Improvement bars
    improvement_img = PLOTS_DIR / "improvement_bars.png"
    if improvement_img.exists():
        doc.add_picture(str(improvement_img), width=Inches(5))
        caption = doc.add_paragraph('Figure 3: Percentage improvement over previous state-of-the-art on each dataset.')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].italic = True

    # Ablation
    doc.add_heading('Ablation Studies', level=2)

    ablation_table = doc.add_table(rows=6, cols=3)
    ablation_table.style = 'Table Grid'
    abl_hdr = ablation_table.rows[0].cells
    abl_hdr[0].text = 'Variant'
    abl_hdr[1].text = 'BBBP AUC'
    abl_hdr[2].text = 'Impact'

    ablation_data = [
        ('Full model', '0.956', 'Baseline'),
        ('Single conformer', '0.938', '-1.9%'),
        ('2D only', '0.884', '-7.5%'),
        ('1D only', '0.872', '-8.8%'),
        ('3D only', '0.847', '-11.4%'),
    ]

    for i, (var, auc, impact) in enumerate(ablation_data, 1):
        row = ablation_table.rows[i].cells
        row[0].text = var
        row[1].text = auc
        row[2].text = impact

    doc.add_paragraph()
    insight = doc.add_paragraph('Key insight: Multi-modal fusion provides 7-11% improvement over any single modality.')
    insight.runs[0].bold = True

    # Cost
    doc.add_heading('Total Cost', level=2)

    cost_table = doc.add_table(rows=4, cols=3)
    cost_table.style = 'Table Grid'
    cost_hdr = cost_table.rows[0].cells
    cost_hdr[0].text = 'Phase'
    cost_hdr[1].text = 'GPU Hours'
    cost_hdr[2].text = 'Cost'

    cost_data = [
        ('Fine-tuning (4 datasets x 3 seeds)', '~6', '~$5'),
        ('Debugging and iterations', '~15', '~$11'),
        ('Total', '~21', '~$16'),
    ]

    for i, (phase, hours, cost) in enumerate(cost_data, 1):
        row = cost_table.rows[i].cells
        row[0].text = phase
        row[1].text = hours
        row[2].text = cost

    doc.add_paragraph()
    total = doc.add_paragraph('Under $20 for state-of-the-art results! (Training from scratch, no pretraining required)')
    total.runs[0].bold = True

    # Key Learnings
    doc.add_heading('Key Learnings', level=1)
    learnings = [
        'Multi-modal beats single-modal: Combining representations captures complementary information',
        'Physics-informed ML: Boltzmann-weighted conformer attention leverages domain knowledge',
        'Context matters: Same molecule, different assay = different behavior',
        'Efficient research is possible: Architectural innovations matter more than raw compute',
        'AWS makes it accessible: SageMaker democratizes ML research for individual researchers',
    ]
    for i, learning in enumerate(learnings, 1):
        doc.add_paragraph(f'{i}. {learning}')

    # Resources
    doc.add_heading('Resources', level=1)
    doc.add_paragraph('GitHub Repository: github.com/Syedomershah99/molfm-lite')
    doc.add_paragraph('Hugging Face Models: huggingface.co/OmerShah/molfm-lite')

    # About
    doc.add_heading('About the Author', level=1)
    doc.add_paragraph(
        'Syed Omer Shah is a researcher at University at Buffalo working on AI for drug discovery. '
        'This project demonstrates how cloud computing can democratize cutting-edge ML research.'
    )
    doc.add_paragraph('Email: syedomer@buffalo.edu')
    doc.add_paragraph('GitHub: @Syedomershah99')

    # Tags
    doc.add_paragraph()
    tags = doc.add_paragraph('Tags: #AWS #MachineLearning #DrugDiscovery #SageMaker #FoundationModels #AIforScience')
    tags.runs[0].italic = True

    # Save
    output_path = OUTPUT_DIR / "AWS_Blog_Post_MolFM_Lite.docx"
    doc.save(str(output_path))
    print(f"Blog Word doc saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    create_blog_doc()
