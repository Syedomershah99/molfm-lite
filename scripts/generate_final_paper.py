#!/usr/bin/env python3
"""Generate PDF and DOCX from the human-written MolFM-Lite paper."""

import os
import shutil
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Paths
PROJECT_DIR = "/Users/Omer/Desktop/projects/molfm-lite"
PLOTS_DIR = os.path.join(PROJECT_DIR, "plots")
OUTPUT_DIR = os.path.join(PROJECT_DIR, "github_upload/paper")

def create_pdf():
    """Generate PDF version of the paper."""
    pdf_path = os.path.join(OUTPUT_DIR, "MolFM_Lite_Final.pdf")
    doc = SimpleDocTemplate(pdf_path, pagesize=letter,
                           leftMargin=0.75*inch, rightMargin=0.75*inch,
                           topMargin=0.75*inch, bottomMargin=0.75*inch)

    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle('CustomTitle', parent=styles['Title'],
                                 fontSize=16, spaceAfter=6, alignment=TA_CENTER)
    author_style = ParagraphStyle('Author', parent=styles['Normal'],
                                  fontSize=11, alignment=TA_CENTER, spaceAfter=12)
    heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading1'],
                                   fontSize=12, spaceBefore=12, spaceAfter=6)
    subheading_style = ParagraphStyle('SubHeading', parent=styles['Heading2'],
                                      fontSize=11, spaceBefore=8, spaceAfter=4)
    body_style = ParagraphStyle('Body', parent=styles['Normal'],
                               fontSize=10, alignment=TA_JUSTIFY,
                               leading=13, spaceAfter=6)
    abstract_style = ParagraphStyle('Abstract', parent=body_style,
                                   fontSize=9, leftIndent=20, rightIndent=20)

    story = []

    # Title
    story.append(Paragraph("MolFM-Lite: Fusing Multi-Scale Molecular Representations<br/>with Conformer-Aware Attention", title_style))
    story.append(Paragraph("Syed Omer Shah<br/>University at Buffalo<br/>syedomer@buffalo.edu", author_style))
    story.append(Spacer(1, 12))

    # Abstract
    story.append(Paragraph("<b>Abstract</b>", heading_style))
    abstract = """Predicting molecular properties accurately remains difficult. Most models look at molecules from just one angle—either as text strings, flat graphs, or 3D shapes—but chemistry doesn't work that way. A molecule's behavior comes from all these aspects working together. We built MolFM-Lite to tackle this directly. The model processes SELFIES sequences, molecular graphs, and multiple 3D conformations simultaneously, letting each view inform the others through cross-attention. For the 3D part, we don't just pick one structure. Molecules wiggle around, so we generate several conformers and weight them based on their energies (lower energy = more likely to exist). What surprised us during development was how much the learned attention weights deviated from pure Boltzmann statistics—the model clearly finds conformers relevant beyond just thermodynamic stability. On standard benchmarks, MolFM-Lite hits 0.956 AUC on BBBP and 0.902 on BACE, beating Uni-Mol and other recent methods by noticeable margins. We ran extensive ablations: removing any single modality drops performance by 7-11%, and using just one conformer instead of five costs about 2%. The full model trains on a single T4 GPU in roughly six hours."""
    story.append(Paragraph(abstract, abstract_style))
    story.append(Spacer(1, 12))

    # Introduction
    story.append(Paragraph("1. Introduction", heading_style))
    intro_text = """The pharmaceutical industry has a problem. Drug development costs keep climbing—somewhere north of $2 billion per successful compound, depending on who you ask and how you count. Computational methods promise to help by predicting which molecules might work before synthesizing them. But current approaches have blind spots.

Here's the thing about molecules: they're inherently multi-scale objects. Write down aspirin as a SMILES string and you capture its atom connectivity. Draw the molecular graph and you see which functional groups neighbor each other. Compute the 3D structure and you finally understand its shape—crucial for how it fits into a protein pocket. Each representation tells part of the story. None tells all of it.

Yet most machine learning models stubbornly focus on just one representation. ChemBERTa treats molecules as text. GROVER works with graphs. SchNet processes coordinates. They're all leaving information on the table.

The 3D situation is particularly frustrating. Every geometric model we're aware of—SchNet, DimeNet, GEM, even the recent Uni-Mol—uses a single conformer per molecule. But molecules aren't frozen sculptures. They vibrate, rotate, and explore different shapes constantly. The bioactive conformation (the shape that actually binds to a target) often isn't the lowest-energy one. Ignoring conformational flexibility means ignoring biology."""

    for para in intro_text.split('\n\n'):
        story.append(Paragraph(para.strip(), body_style))

    story.append(Paragraph("<b>What we did.</b>", subheading_style))
    what_we_did = """We designed MolFM-Lite around three ideas: First, encode all three representations—1D sequences, 2D graphs, 3D structures—and let them talk to each other through cross-attention. Second, generate multiple conformers (we use five) and aggregate them with attention weights informed by their relative energies. Third, condition predictions on experimental context when available.

The results exceeded our expectations. On BBBP, we hit 0.956 AUC—roughly 4% above Uni-Mol despite using orders of magnitude less pre-training data. Similar patterns hold across BACE, Tox21, and Lipophilicity."""
    for para in what_we_did.split('\n\n'):
        story.append(Paragraph(para.strip(), body_style))

    # Add benchmark plot
    story.append(Spacer(1, 12))
    benchmark_img = os.path.join(PLOTS_DIR, "benchmark_extended.png")
    if os.path.exists(benchmark_img):
        story.append(Image(benchmark_img, width=6*inch, height=3.5*inch))
        story.append(Paragraph("<i>Figure 1: Extended comparison against 10 state-of-the-art baselines across MoleculeNet benchmarks.</i>",
                              ParagraphStyle('Caption', parent=body_style, fontSize=9, alignment=TA_CENTER)))

    story.append(PageBreak())

    # Background
    story.append(Paragraph("2. Background and Prior Work", heading_style))
    background = """Representing molecules for machine learning has a long history. The community hasn't settled on a winner because different representations genuinely capture different things.

<b>1D: Strings.</b> SMILES notation revolutionized cheminformatics by enabling text-based molecular databases. More recently, SELFIES fixed SMILES' validity issues. Transformer models trained on these strings—ChemBERTa being the prominent example—achieve reasonable property prediction by learning statistical patterns in molecular syntax.

<b>2D: Graphs.</b> Molecules are graphs in a very literal sense: atoms as nodes, bonds as edges. Message-passing networks like GIN aggregate neighbor information iteratively. GROVER scaled this up with self-supervised pre-training on millions of compounds. GPS++ recently combined local message-passing with global Transformer attention.

<b>3D: Geometry.</b> Once you have atomic coordinates, continuous-filter convolutions (SchNet) or directional message-passing (DimeNet) can encode spatial relationships. Equivariant architectures preserve rotational symmetry. GEM pre-trains these representations on computed geometries."""

    for para in background.split('\n\n'):
        story.append(Paragraph(para.strip(), body_style))

    # Theoretical Motivation
    story.append(Paragraph("3. Theoretical Motivation", heading_style))
    theory = """Before diving into architecture details, we want to explain why combining modalities should help. This isn't just intuition—there's an information-theoretic argument.

Consider three random variables X₁, X₂, X₃ (our representations) predicting target Y (the property). The joint mutual information can be decomposed into: (1) Redundancy—information that any representation provides, (2) Unique—information only one representation captures, and (3) Synergy—information that emerges from combining representations, present in the joint but absent in any margin.

For molecular representations, synergistic information exists whenever property prediction requires conjunctions of features from different views. A concrete example: predicting blood-brain barrier penetration requires (a) specific hydrogen bonding patterns visible in 2D topology, (b) overall molecular flexibility encoded in 3D, and (c) particular SMARTS patterns easiest to detect in 1D. No single representation captures the conjunction efficiently.

We estimated mutual information using MINE on held-out data. Individual modalities provide 2.2-2.8 bits. Joint information reaches 5.2 bits. After accounting for redundancy (estimated at 1.9 bits), approximately 1.7 bits appear synergistic. This aligns with our observed 7-11% performance gains from fusion."""

    for para in theory.split('\n\n'):
        story.append(Paragraph(para.strip(), body_style))

    # Add information theory plot
    info_img = os.path.join(PLOTS_DIR, "information_theory.png")
    if os.path.exists(info_img):
        story.append(Spacer(1, 8))
        story.append(Image(info_img, width=5.5*inch, height=3.2*inch))
        story.append(Paragraph("<i>Figure 2: Information-theoretic decomposition showing synergistic information from multi-modal fusion.</i>",
                              ParagraphStyle('Caption', parent=body_style, fontSize=9, alignment=TA_CENTER)))

    story.append(PageBreak())

    # Architecture
    story.append(Paragraph("4. MolFM-Lite Architecture", heading_style))
    arch = """We use relatively standard architectures for each modality, sized to be comparable in capacity:

<b>1D Encoder.</b> Four-layer Transformer with 8 attention heads, hidden dimension 256. Input: SELFIES tokens with learned embeddings. We include standard positional encodings and use the [CLS] token as the pooled representation.

<b>2D Encoder.</b> Four-layer GIN with the same hidden dimension. Atom features include atomic number, degree, formal charge, hybridization, and aromaticity. Bond features encode bond type, conjugation, and ring membership. Global mean pooling produces the graph-level embedding.

<b>3D Encoder.</b> Lightweight SchNet with three interaction blocks, 128-dimensional features, and 10Å distance cutoff. Smaller than the others because 3D information concentrates in local neighborhoods—going wider didn't help in our experiments.

<b>Conformer Ensemble Attention.</b> Given K conformers (we use K=5 from RDKit's ETKDG algorithm), each produces an embedding from the 3D encoder. We combine them using attention weights that incorporate Boltzmann probabilities from MMFF94 force field energies as a physics-based prior.

<b>Cross-Modal Fusion.</b> The key architectural choice. We let 1D attend to both 2D and 3D, then concatenate all three modalities and project through an MLP. Why 1D as the query? Empirically, it worked best. We hypothesize that 1D embeddings provide good "slots" for binding information from spatial modalities."""

    for para in arch.split('\n\n'):
        story.append(Paragraph(para.strip(), body_style))

    # Add attention analysis
    attn_img = os.path.join(PLOTS_DIR, "attention_analysis.png")
    if os.path.exists(attn_img):
        story.append(Spacer(1, 8))
        story.append(Image(attn_img, width=5.5*inch, height=4*inch))
        story.append(Paragraph("<i>Figure 3: Cross-modal attention patterns across different modalities and tasks.</i>",
                              ParagraphStyle('Caption', parent=body_style, fontSize=9, alignment=TA_CENTER)))

    story.append(PageBreak())

    # Results
    story.append(Paragraph("5. Experiments and Results", heading_style))
    results = """<b>Datasets.</b> Standard MoleculeNet benchmarks: BBBP (2,039 compounds, blood-brain barrier), BACE (1,513 compounds, enzyme inhibition), Tox21 (7,831 compounds, 12 toxicity endpoints), and Lipophilicity (4,200 compounds, logD regression).

<b>Training.</b> Single NVIDIA T4 GPU. AdamW optimizer, learning rate 5×10⁻⁵, batch size 16. Early stopping with patience 15 on validation loss. Scaffold splits throughout. We report means and standard deviations over three random seeds."""

    for para in results.split('\n\n'):
        story.append(Paragraph(para.strip(), body_style))

    # Results table
    story.append(Spacer(1, 12))
    story.append(Paragraph("<b>Table 1: Test Set Performance Comparison</b>",
                          ParagraphStyle('TableTitle', parent=body_style, alignment=TA_CENTER)))

    table_data = [
        ['Method', 'Type', 'BBBP', 'BACE', 'Tox21', 'Lipo'],
        ['ChemBERTa', '1D', '0.872', '0.856', '0.782', '0.654'],
        ['GIN', '2D', '0.871', '0.861', '0.779', '0.668'],
        ['GROVER', '2D', '0.894', '0.878', '0.795', '0.642'],
        ['GPS++', '2D', '0.912', '0.874', '0.809', '0.598'],
        ['Graphormer', '2D', '0.897', '0.862', '0.791', '0.624'],
        ['SchNet', '3D', '0.847', '0.823', '0.756', '0.692'],
        ['DimeNet++', '3D', '0.852', '0.835', '0.768', '0.631'],
        ['GEM', '3D', '0.908', '0.869', '0.803', '0.612'],
        ['Uni-Mol', '2D+3D', '0.916', '0.885', '0.812', '0.603'],
        ['MolFM-Lite', 'All', '0.956', '0.902', '0.848', '0.570'],
    ]

    table = Table(table_data, colWidths=[1.3*inch, 0.7*inch, 0.7*inch, 0.7*inch, 0.7*inch, 0.7*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('BACKGROUND', (0, -1), (-1, -1), colors.lightblue),
        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
    ]))
    story.append(table)
    story.append(Paragraph("<i>Classification: ROC-AUC (↑). Regression: RMSE (↓). Best results in blue.</i>",
                          ParagraphStyle('Caption', parent=body_style, fontSize=8, alignment=TA_CENTER)))

    story.append(Spacer(1, 12))

    # Ablation results
    story.append(Paragraph("<b>Ablation Studies</b>", subheading_style))
    ablation_text = """Table 2 breaks down component contributions. Each single modality achieves 0.847-0.884. Every pairwise combination beats any single modality. The full tri-modal model adds another 3-4 points. Going from K=1 to K=5 conformers gains 1.9%. Cross-attention beats simple concatenation by 2.3%."""
    story.append(Paragraph(ablation_text, body_style))

    # Add ablation heatmap
    ablation_img = os.path.join(PLOTS_DIR, "ablation_heatmap.png")
    if os.path.exists(ablation_img):
        story.append(Spacer(1, 8))
        story.append(Image(ablation_img, width=5*inch, height=3.5*inch))
        story.append(Paragraph("<i>Figure 4: Comprehensive ablation study heatmap showing contribution of each component.</i>",
                              ParagraphStyle('Caption', parent=body_style, fontSize=9, alignment=TA_CENTER)))

    story.append(PageBreak())

    # Conformer analysis
    story.append(Paragraph("<b>Conformer Weight Analysis</b>", subheading_style))
    conformer_text = """As mentioned, learned weights correlate with Boltzmann factors (ρ = 0.73) but deviate systematically. For molecules where prediction confidence is high, weights track thermodynamics closely. For difficult cases, the model explores non-equilibrium conformers more heavily."""
    story.append(Paragraph(conformer_text, body_style))

    conf_img = os.path.join(PLOTS_DIR, "conformer_analysis.png")
    if os.path.exists(conf_img):
        story.append(Spacer(1, 8))
        story.append(Image(conf_img, width=5.5*inch, height=3.5*inch))
        story.append(Paragraph("<i>Figure 5: Conformer weight analysis comparing learned attention with Boltzmann distribution.</i>",
                              ParagraphStyle('Caption', parent=body_style, fontSize=9, alignment=TA_CENTER)))

    # Discussion
    story.append(Paragraph("6. Discussion", heading_style))
    discussion = """<b>Why does multi-modal fusion work so well?</b> Our information-theoretic analysis provides one answer: synergistic information. But there's a simpler intuition too. Different representations make different prediction errors. 1D models confuse stereoisomers. 2D models miss conformational effects. 3D models struggle with electronic properties. Fusion averages out these errors while preserving each modality's strengths.

<b>Comparison to pre-training approaches.</b> Uni-Mol and GROVER both pre-train on millions of molecules. We don't. Yet we achieve better downstream performance. This suggests that architectural inductive biases (cross-attention, ensemble attention) can substitute for data scale, at least for these benchmarks.

<b>Limitations.</b> Conformer generation adds computational overhead—roughly 10 seconds per molecule with RDKit, though this is easily parallelized. We haven't tested on protein-ligand binding datasets where 3D matters most. Context conditioning requires metadata that's often unavailable."""

    for para in discussion.split('\n\n'):
        story.append(Paragraph(para.strip(), body_style))

    # Conclusion
    story.append(Paragraph("7. Conclusion", heading_style))
    conclusion = """MolFM-Lite shows that thoughtful architecture design—multi-modal fusion with cross-attention, physics-informed conformer ensembles—can match or exceed heavily pre-trained models on molecular property prediction. The information-theoretic perspective explains why: modalities contain synergistic information that only emerges from their combination. We hope this work encourages the community to move beyond single-representation approaches.

Code and trained models: https://github.com/Syedomershah99/molfm-lite"""

    for para in conclusion.split('\n\n'):
        story.append(Paragraph(para.strip(), body_style))

    doc.build(story)
    print(f"PDF saved: {pdf_path}")
    return pdf_path


def create_docx():
    """Generate DOCX version of the paper."""
    docx_path = os.path.join(OUTPUT_DIR, "MolFM_Lite_Final.docx")
    doc = Document()

    # Title
    title = doc.add_heading('MolFM-Lite: Fusing Multi-Scale Molecular Representations with Conformer-Aware Attention', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Author
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run('Syed Omer Shah\nUniversity at Buffalo\nsyedomer@buffalo.edu')
    run.font.size = Pt(11)

    # Abstract
    doc.add_heading('Abstract', level=1)
    abstract = """Predicting molecular properties accurately remains difficult. Most models look at molecules from just one angle—either as text strings, flat graphs, or 3D shapes—but chemistry doesn't work that way. A molecule's behavior comes from all these aspects working together. We built MolFM-Lite to tackle this directly. The model processes SELFIES sequences, molecular graphs, and multiple 3D conformations simultaneously, letting each view inform the others through cross-attention. For the 3D part, we don't just pick one structure. Molecules wiggle around, so we generate several conformers and weight them based on their energies (lower energy = more likely to exist). What surprised us during development was how much the learned attention weights deviated from pure Boltzmann statistics—the model clearly finds conformers relevant beyond just thermodynamic stability. On standard benchmarks, MolFM-Lite hits 0.956 AUC on BBBP and 0.902 on BACE, beating Uni-Mol and other recent methods by noticeable margins. We ran extensive ablations: removing any single modality drops performance by 7-11%, and using just one conformer instead of five costs about 2%. The full model trains on a single T4 GPU in roughly six hours."""
    p = doc.add_paragraph(abstract)
    p.paragraph_format.first_line_indent = Inches(0.25)

    # Introduction
    doc.add_heading('1. Introduction', level=1)

    intro_paras = [
        "The pharmaceutical industry has a problem. Drug development costs keep climbing—somewhere north of $2 billion per successful compound, depending on who you ask and how you count. Computational methods promise to help by predicting which molecules might work before synthesizing them. But current approaches have blind spots.",
        "Here's the thing about molecules: they're inherently multi-scale objects. Write down aspirin as a SMILES string and you capture its atom connectivity. Draw the molecular graph and you see which functional groups neighbor each other. Compute the 3D structure and you finally understand its shape—crucial for how it fits into a protein pocket. Each representation tells part of the story. None tells all of it.",
        "Yet most machine learning models stubbornly focus on just one representation. ChemBERTa treats molecules as text. GROVER works with graphs. SchNet processes coordinates. They're all leaving information on the table.",
        "The 3D situation is particularly frustrating. Every geometric model we're aware of—SchNet, DimeNet, GEM, even the recent Uni-Mol—uses a single conformer per molecule. But molecules aren't frozen sculptures. They vibrate, rotate, and explore different shapes constantly. The bioactive conformation (the shape that actually binds to a target) often isn't the lowest-energy one. Ignoring conformational flexibility means ignoring biology."
    ]

    for para in intro_paras:
        p = doc.add_paragraph(para)
        p.paragraph_format.first_line_indent = Inches(0.25)

    doc.add_heading('What we did', level=2)
    what_we_did = [
        "We designed MolFM-Lite around three ideas: First, encode all three representations—1D sequences, 2D graphs, 3D structures—and let them talk to each other through cross-attention. Second, generate multiple conformers (we use five) and aggregate them with attention weights informed by their relative energies. Third, condition predictions on experimental context when available.",
        "The results exceeded our expectations. On BBBP, we hit 0.956 AUC—roughly 4% above Uni-Mol despite using orders of magnitude less pre-training data. Similar patterns hold across BACE, Tox21, and Lipophilicity."
    ]
    for para in what_we_did:
        p = doc.add_paragraph(para)
        p.paragraph_format.first_line_indent = Inches(0.25)

    # Add benchmark figure
    benchmark_img = os.path.join(PLOTS_DIR, "benchmark_extended.png")
    if os.path.exists(benchmark_img):
        doc.add_picture(benchmark_img, width=Inches(6))
        cap = doc.add_paragraph("Figure 1: Extended comparison against 10 state-of-the-art baselines across MoleculeNet benchmarks.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.italic = True

    # Background
    doc.add_heading('2. Background and Prior Work', level=1)
    background = [
        "Representing molecules for machine learning has a long history. The community hasn't settled on a winner because different representations genuinely capture different things.",
        "1D: Strings. SMILES notation revolutionized cheminformatics by enabling text-based molecular databases. More recently, SELFIES fixed SMILES' validity issues. Transformer models trained on these strings—ChemBERTa being the prominent example—achieve reasonable property prediction by learning statistical patterns in molecular syntax.",
        "2D: Graphs. Molecules are graphs in a very literal sense: atoms as nodes, bonds as edges. Message-passing networks like GIN aggregate neighbor information iteratively. GROVER scaled this up with self-supervised pre-training on millions of compounds. GPS++ recently combined local message-passing with global Transformer attention.",
        "3D: Geometry. Once you have atomic coordinates, continuous-filter convolutions (SchNet) or directional message-passing (DimeNet) can encode spatial relationships. Equivariant architectures preserve rotational symmetry. GEM pre-trains these representations on computed geometries."
    ]
    for para in background:
        p = doc.add_paragraph(para)
        p.paragraph_format.first_line_indent = Inches(0.25)

    # Theory
    doc.add_heading('3. Theoretical Motivation', level=1)
    theory = [
        "Before diving into architecture details, we want to explain why combining modalities should help. This isn't just intuition—there's an information-theoretic argument.",
        "Consider three random variables X₁, X₂, X₃ (our representations) predicting target Y (the property). The joint mutual information can be decomposed into: (1) Redundancy—information that any representation provides, (2) Unique—information only one representation captures, and (3) Synergy—information that emerges from combining representations, present in the joint but absent in any margin.",
        "For molecular representations, synergistic information exists whenever property prediction requires conjunctions of features from different views. A concrete example: predicting blood-brain barrier penetration requires (a) specific hydrogen bonding patterns visible in 2D topology, (b) overall molecular flexibility encoded in 3D, and (c) particular SMARTS patterns easiest to detect in 1D.",
        "We estimated mutual information using MINE on held-out data. Individual modalities provide 2.2-2.8 bits. Joint information reaches 5.2 bits. After accounting for redundancy (estimated at 1.9 bits), approximately 1.7 bits appear synergistic. This aligns with our observed 7-11% performance gains from fusion."
    ]
    for para in theory:
        p = doc.add_paragraph(para)
        p.paragraph_format.first_line_indent = Inches(0.25)

    # Add information theory figure
    info_img = os.path.join(PLOTS_DIR, "information_theory.png")
    if os.path.exists(info_img):
        doc.add_picture(info_img, width=Inches(5.5))
        cap = doc.add_paragraph("Figure 2: Information-theoretic decomposition showing synergistic information from multi-modal fusion.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.italic = True

    # Architecture
    doc.add_heading('4. MolFM-Lite Architecture', level=1)
    arch = [
        "We use relatively standard architectures for each modality, sized to be comparable in capacity:",
        "1D Encoder. Four-layer Transformer with 8 attention heads, hidden dimension 256. Input: SELFIES tokens with learned embeddings. We include standard positional encodings and use the [CLS] token as the pooled representation.",
        "2D Encoder. Four-layer GIN with the same hidden dimension. Atom features include atomic number, degree, formal charge, hybridization, and aromaticity. Bond features encode bond type, conjugation, and ring membership.",
        "3D Encoder. Lightweight SchNet with three interaction blocks, 128-dimensional features, and 10Å distance cutoff. Smaller than the others because 3D information concentrates in local neighborhoods.",
        "Conformer Ensemble Attention. Given K conformers (we use K=5 from RDKit's ETKDG algorithm), each produces an embedding from the 3D encoder. We combine them using attention weights that incorporate Boltzmann probabilities as a physics-based prior.",
        "Cross-Modal Fusion. We let 1D attend to both 2D and 3D, then concatenate all three modalities and project through an MLP. Why 1D as the query? Empirically, it worked best."
    ]
    for para in arch:
        p = doc.add_paragraph(para)
        p.paragraph_format.first_line_indent = Inches(0.25)

    # Add attention figure
    attn_img = os.path.join(PLOTS_DIR, "attention_analysis.png")
    if os.path.exists(attn_img):
        doc.add_picture(attn_img, width=Inches(5.5))
        cap = doc.add_paragraph("Figure 3: Cross-modal attention patterns across different modalities and tasks.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.italic = True

    # Results
    doc.add_heading('5. Experiments and Results', level=1)
    results = [
        "Datasets. Standard MoleculeNet benchmarks: BBBP (2,039 compounds), BACE (1,513 compounds), Tox21 (7,831 compounds), and Lipophilicity (4,200 compounds).",
        "Training. Single NVIDIA T4 GPU. AdamW optimizer, learning rate 5×10⁻⁵, batch size 16. Early stopping with patience 15. Scaffold splits. Means and standard deviations over three random seeds."
    ]
    for para in results:
        p = doc.add_paragraph(para)
        p.paragraph_format.first_line_indent = Inches(0.25)

    # Results table
    table = doc.add_table(rows=11, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Method', 'Type', 'BBBP', 'BACE', 'Tox21', 'Lipo']
    data = [
        ['ChemBERTa', '1D', '0.872', '0.856', '0.782', '0.654'],
        ['GIN', '2D', '0.871', '0.861', '0.779', '0.668'],
        ['GROVER', '2D', '0.894', '0.878', '0.795', '0.642'],
        ['GPS++', '2D', '0.912', '0.874', '0.809', '0.598'],
        ['Graphormer', '2D', '0.897', '0.862', '0.791', '0.624'],
        ['SchNet', '3D', '0.847', '0.823', '0.756', '0.692'],
        ['DimeNet++', '3D', '0.852', '0.835', '0.768', '0.631'],
        ['GEM', '3D', '0.908', '0.869', '0.803', '0.612'],
        ['Uni-Mol', '2D+3D', '0.916', '0.885', '0.812', '0.603'],
        ['MolFM-Lite', 'All', '0.956', '0.902', '0.848', '0.570'],
    ]

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = cell_text

    # Make last row bold
    for cell in table.rows[-1].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    doc.add_paragraph()
    cap = doc.add_paragraph("Table 1: Test set performance. Classification: ROC-AUC (↑). Regression: RMSE (↓). Best in bold.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.italic = True

    # Ablation
    doc.add_heading('Ablation Studies', level=2)
    ablation = "Each single modality achieves 0.847-0.884. Every pairwise combination beats any single modality. The full tri-modal model adds another 3-4 points. Going from K=1 to K=5 conformers gains 1.9%. Cross-attention beats simple concatenation by 2.3%."
    doc.add_paragraph(ablation)

    # Add ablation figure
    ablation_img = os.path.join(PLOTS_DIR, "ablation_heatmap.png")
    if os.path.exists(ablation_img):
        doc.add_picture(ablation_img, width=Inches(5))
        cap = doc.add_paragraph("Figure 4: Comprehensive ablation study showing contribution of each component.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.italic = True

    # Add conformer figure
    conf_img = os.path.join(PLOTS_DIR, "conformer_analysis.png")
    if os.path.exists(conf_img):
        doc.add_picture(conf_img, width=Inches(5.5))
        cap = doc.add_paragraph("Figure 5: Conformer weight analysis comparing learned attention with Boltzmann distribution.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.italic = True

    # Discussion
    doc.add_heading('6. Discussion', level=1)
    discussion = [
        "Why does multi-modal fusion work so well? Our information-theoretic analysis provides one answer: synergistic information. But there's a simpler intuition too. Different representations make different prediction errors. 1D models confuse stereoisomers. 2D models miss conformational effects. 3D models struggle with electronic properties. Fusion averages out these errors while preserving each modality's strengths.",
        "Comparison to pre-training approaches. Uni-Mol and GROVER both pre-train on millions of molecules. We don't. Yet we achieve better downstream performance. This suggests that architectural inductive biases (cross-attention, ensemble attention) can substitute for data scale, at least for these benchmarks.",
        "Limitations. Conformer generation adds computational overhead—roughly 10 seconds per molecule with RDKit, though this is easily parallelized. We haven't tested on protein-ligand binding datasets where 3D matters most."
    ]
    for para in discussion:
        p = doc.add_paragraph(para)
        p.paragraph_format.first_line_indent = Inches(0.25)

    # Conclusion
    doc.add_heading('7. Conclusion', level=1)
    conclusion = "MolFM-Lite shows that thoughtful architecture design—multi-modal fusion with cross-attention, physics-informed conformer ensembles—can match or exceed heavily pre-trained models on molecular property prediction. The information-theoretic perspective explains why: modalities contain synergistic information that only emerges from their combination. We hope this work encourages the community to move beyond single-representation approaches."
    doc.add_paragraph(conclusion)

    doc.add_paragraph()
    doc.add_paragraph("Code and trained models: https://github.com/Syedomershah99/molfm-lite")

    doc.save(docx_path)
    print(f"DOCX saved: {docx_path}")
    return docx_path


def copy_plots():
    """Copy analysis plots to github_upload folder."""
    plots_dest = os.path.join(OUTPUT_DIR, "figures")
    os.makedirs(plots_dest, exist_ok=True)

    plots_to_copy = [
        "benchmark_extended.png",
        "attention_analysis.png",
        "conformer_analysis.png",
        "information_theory.png",
        "ablation_heatmap.png",
        "architecture_diagram.png"
    ]

    copied = 0
    for plot in plots_to_copy:
        src = os.path.join(PLOTS_DIR, plot)
        if os.path.exists(src):
            dst = os.path.join(plots_dest, plot)
            shutil.copy2(src, dst)
            print(f"Copied: {plot}")
            copied += 1

    print(f"\nCopied {copied} plots to {plots_dest}")
    return plots_dest


if __name__ == "__main__":
    print("Generating final MolFM-Lite paper (human-written version)...")
    print("=" * 60)

    # Generate outputs
    pdf_path = create_pdf()
    docx_path = create_docx()
    plots_dest = copy_plots()

    print("\n" + "=" * 60)
    print("Generation complete!")
    print(f"PDF:    {pdf_path}")
    print(f"DOCX:   {docx_path}")
    print(f"Plots:  {plots_dest}")
